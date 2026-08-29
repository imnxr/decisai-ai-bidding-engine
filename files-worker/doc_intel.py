"""
DOCUMENT INTELLIGENCE — the hackathon spec's actual "Module 1".

PDF/DOCX/TXT -> raw text -> chunks -> structured RFP JSON with
requirements, evaluation criteria, deadlines, questions, financials.

WHY heuristic (regex) extraction, not LLM extraction: a small local model doing
free-form extraction live on stage is the highest-risk demo move possible.
Regex rules are deterministic, instant, fully offline, and explainable to
judges line by line. The LLM is reserved for what it's good at (drafting
prose), not for structured extraction.
"""
import os
import re
import uuid

# ---------------- text extraction (lazy imports so .txt works dep-free) ----------------

def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        try:
            import fitz  # PyMuPDF — primary (fast, good layout handling)
            with fitz.open(path) as doc:
                return "\n".join(page.get_text() for page in doc)
        except ImportError:
            # graceful fallback so a missing pymupdf doesn't kill the demo
            from pdfminer.high_level import extract_text as pdfminer_extract
            return pdfminer_extract(path)
    if ext == ".docx":
        import docx  # python-docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:  # RFPs love tables — don't lose them
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext in (".txt", ".md"):
        with open(path, encoding="utf-8", errors="replace") as f:
            return f.read()
    raise ValueError(f"Unsupported file type: {ext} (use .pdf, .docx, or .txt)")


def chunk_text(text: str, max_chars: int = 1200, overlap: int = 150) -> list:
    """Simple sliding-window chunks for downstream RAG over the RFP itself."""
    chunks, start, n = [], 0, len(text)
    while start < n:
        end = min(start + max_chars, n)
        if end < n:  # try to break on a sentence/line boundary
            cut = max(text.rfind(". ", start, end), text.rfind("\n", start, end))
            if cut > start + max_chars // 2:
                end = cut + 1
        chunks.append({"chunk_id": f"CH-{len(chunks)+1:03d}",
                       "text": text[start:end].strip()})
        if end >= n:
            break
        start = end - overlap
    return chunks


# ---------------- extraction rules ----------------

MANDATORY_RE = re.compile(r"\b(shall|must|is required|are required|mandatory)\b", re.I)
OPTIONAL_RE = re.compile(r"\b(should|may|preferred|desirable|advantageous)\b", re.I)

# Some mandatory-looking clauses are not capability questions. Matching "bid
# security" or "retention" against past projects creates fake PASS rows, so we
# keep them visible but mark them as administrative information.
ADMIN_RE = re.compile(
    r"\b(bid security|performance security|earnest money|bank guarantee|"
    r"retention|payment|payments|tax|taxes|proposal validity|validity period|"
    r"submission deadline|pre-bid|prebid|sealed proposals|tender fee|"
    r"must accompany each proposal)\b", re.I)

# Map requirement keywords -> bid Sector, so retrieve_evidence() gets its
# category boost (bridges RFP language to the capability library domains).
CATEGORY_KEYWORDS = [
    # NOTE: deliberately no bare "security" — "bid security" (a financial term)
    # was matching it and miscategorizing financial clauses as IT Services.
    (r"iso 27001|cyber|information security|security controls|security operations|"
     r"firewall|penetration|network|cloud|erp|software|data cent", "IT Services"),
    (r"fleet|logistics|transport|warehous", "Logistics"),
    (r"road|bridge|construction|civil works|pavement", "Construction"),
    (r"hospital|medical|clinic|health", "Healthcare"),
    (r"solar|energy|power plant|grid", "Energy"),
    (r"lms|e-?learning|education|training platform|curriculum", "Education"),
    (r"bank|finance|payment|fintech", "Finance"),
    (r"telecom|fiber|5g|bts", "Telecom"),
]

DATE_RE = re.compile(
    r"(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}"
    r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}"
    r"|\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4})", re.I)

MONEY_RE = re.compile(r"(PKR|Rs\.?|USD|\$)\s?[\d,]+(?:\.\d+)?\s?(?:M|million|billion|B)?", re.I)
PCT_RE = re.compile(r"\b(\d{1,3})\s?%")


def _sentences(text: str) -> list:
    """Split into sentence-ish units; lines and '.'-terminated sentences both count."""
    out = []
    for line in _join_wrapped_lines(text).split("\n"):
        line = line.strip()
        if not line:
            continue
        out.extend(s.strip() for s in re.split(r"(?<=[.;])\s+", line) if len(s.strip()) > 15)
    return out


def _starts_new_item(line: str) -> bool:
    return bool(re.match(
        r"^(SECTION\b|\d+(?:\.\d+)+\b|Q\d+[.:)]|Question\s+\d+|[A-Za-z].*:\s*\d{1,3}\s?%$)",
        line, re.I))


def _join_wrapped_lines(text: str) -> str:
    """PDF extractors often wrap one requirement across multiple lines.
    Rebuild logical lines so clauses like "ISO" + "27001 controls" stay intact."""
    logical, current = [], ""
    for raw in text.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if not line:
            if current:
                logical.append(current)
                current = ""
            continue
        if not current:
            current = line
            continue
        if _starts_new_item(line) or current.endswith((".", ":", "?", "%")):
            logical.append(current)
            current = line
        else:
            current = f"{current} {line}"
    if current:
        logical.append(current)
    return "\n".join(logical)


def guess_category(sentence: str) -> str:
    low = sentence.lower()
    for pattern, sector in CATEGORY_KEYWORDS:
        if re.search(pattern, low):
            return sector
    return ""


def classify_requirement_type(sentence: str) -> str:
    """Return capability vs administrative so Module 2 matches only real evidence needs."""
    return "administrative" if ADMIN_RE.search(sentence) else "capability"


def _is_heading(s: str) -> bool:
    """Section headings like 'SECTION 2 — MANDATORY REQUIREMENTS' contain trigger
    words (mandatory/preferred) but aren't requirements. Detect by shape:
    starts with SECTION, or letters are >=80% uppercase."""
    if re.match(r"^SECTION\b", s, re.I):
        return True
    letters = [c for c in s if c.isalpha()]
    return bool(letters) and sum(c.isupper() for c in letters) / len(letters) >= 0.8


def extract_requirements(text: str) -> list:
    reqs = []
    for s in _sentences(text):
        if s.endswith("?") or _is_heading(s):  # questions/headings handled separately
            continue
        if MANDATORY_RE.search(s):
            mandatory = True
        elif OPTIONAL_RE.search(s):
            mandatory = False
        else:
            continue
        reqs.append({
            "requirement_id": f"R-{len(reqs)+1:03d}",
            "text": s[:500],
            "mandatory": mandatory,
            "category": guess_category(s),
            "requirement_type": classify_requirement_type(s),
        })
    return reqs


def extract_questions(text: str) -> list:
    qs = []
    for s in _sentences(text):
        if s.endswith("?") or re.match(r"^(Q\d+[.:)]|Question\s+\d+)", s, re.I):
            qs.append({"question_id": f"Q-{len(qs)+1:03d}",
                       "question": s[:500],
                       "category": guess_category(s)})
    return qs


def extract_deadlines(text: str) -> list:
    dls = []
    for s in _sentences(text):
        for m in DATE_RE.finditer(s):
            dls.append({"date": m.group(0), "context": s[:300]})
    return dls


def extract_financials(text: str) -> list:
    fins = []
    for s in _sentences(text):
        for m in MONEY_RE.finditer(s):
            fins.append({"amount": m.group(0).strip(), "context": s[:300]})
    return fins


# A criterion line looks like "Some criterion name: 35%" — name then weight at END
# of line. Anchoring to line end avoids false hits like "with 10% retention until..."
# WHY not section-tracking: PDF extractors insert unpredictable blank lines/ordering,
# so "am I inside the criteria section?" state is fragile. Line shape is not.
CRIT_LINE_RE = re.compile(r"^(?P<name>[A-Za-z][\w\s,&/()'\-]{3,90}?)[:\-–]?\s*(?P<pct>\d{1,3})\s?%\s*$")


def extract_criteria(text: str) -> list:
    crits, seen = [], set()
    for line in text.split("\n"):
        line = line.strip()
        m = CRIT_LINE_RE.match(line)
        if m and line not in seen:
            seen.add(line)
            crits.append({"criterion": line[:300], "weight_pct": int(m.group("pct"))})
    return crits


def extract_meta(text: str) -> dict:
    """Title / issuing org / reference for the RFP summary screen."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    title = next((l for l in lines[:8]
                  if len(l) > 12 and not re.match(r"^request for proposal|^rfp\b", l, re.I)), "")
    m = re.search(r"issuing authority[:\-]?\s*(.+)", text, re.I)
    issuer = m.group(1).strip()[:120] if m else ""
    m = re.search(r"(?:rfp|tender)\s*reference[:\-]?\s*(\S+)", text, re.I)
    return {"title": title[:150], "issuer": issuer, "reference": m.group(1) if m else ""}


# ---------------- optional Claude-powered extraction (hybrid) ----------------

LLM_EXTRACT_PROMPT = """You extract structured data from tender/RFP documents.
Read the document below and output ONLY a JSON object with exactly these keys:
{"requirements": [{"text": "...", "mandatory": true/false, "category": "..."}],
 "evaluation_criteria": [{"criterion": "...", "weight_pct": 0}],
 "questions": [{"question": "...", "category": "..."}]}
Rules: requirements are obligations on the bidder (mandatory if phrased as
shall/must/required, or clearly compulsory in context — including ones NOT using
those words). category must be one of: IT Services, Construction, Healthcare,
Energy, Logistics, Education, Finance, Telecom, or "" if unclear. Copy text
verbatim from the document; do not paraphrase or invent anything.

DOCUMENT:
"""


def llm_enrich(text: str, base: dict, provider: str = "claude") -> dict:
    """
    HYBRID EXTRACTION: ask the selected LLM provider (Claude API = online,
    local Ollama = offline) to extract requirements/criteria/questions — it
    catches obligations phrased without shall/must keywords that regex misses
    on real 50+ page RFPs. The regex output stays as both the fallback and the
    floor: if the LLM result is missing, malformed, or finds FEWER items than
    regex, we keep regex. Deadlines/financials always stay regex (dates and
    money are exactly what deterministic patterns are best at).
    """
    try:
        import json as _json
        if provider == "ollama":
            from ollama_client import call_ollama
            raw = call_ollama(LLM_EXTRACT_PROMPT + text[:12000], timeout=90)
        else:
            from claude_client import call_claude, claude_available
            if not claude_available():
                return base
            raw = call_claude(LLM_EXTRACT_PROMPT + text[:24000], timeout=30, max_tokens=2000)
        if not raw:
            return base
        start, end = raw.find("{"), raw.rfind("}")
        parsed = _json.loads(raw[start:end + 1])

        reqs = []
        for i, r in enumerate(parsed.get("requirements", [])):
            text_value = str(r.get("text", "")).strip()
            if not text_value or _is_heading(text_value):
                continue
            guessed = guess_category(text_value)
            reqs.append({
                "requirement_id": f"R-{i+1:03d}",
                "text": text_value[:500],
                "mandatory": bool(r.get("mandatory", True)),
                "category": guessed or r.get("category", "") or "",
                "requirement_type": classify_requirement_type(text_value),
            })
        crits = [{"criterion": str(c.get("criterion", ""))[:300],
                  "weight_pct": int(c.get("weight_pct", 0))}
                 for c in parsed.get("evaluation_criteria", []) if c.get("criterion")]
        qs = []
        for i, q in enumerate(parsed.get("questions", [])):
            question = str(q.get("question", "")).strip()
            if not question:
                continue
            guessed = guess_category(question)
            qs.append({"question_id": f"Q-{i+1:03d}", "question": question[:500],
                       "category": guessed or q.get("category", "") or ""})

        # LLM must beat regex to replace it (it sees phrasing regex can't)
        if len(reqs) >= len(base["requirements"]):
            base["requirements"] = reqs
        if len(crits) >= len(base["evaluation_criteria"]):
            base["evaluation_criteria"] = crits
        if len(qs) >= len(base["questions"]):
            base["questions"] = qs
        base["extraction_method"] = ("ollama+heuristic" if provider == "ollama"
                                     else "claude+heuristic")
    except Exception:
        pass  # any failure -> regex results stand; never break upload
    return base


# ---------------- main entry ----------------

def parse_rfp(path: str, rfp_id: str = None, use_llm: bool = True,
              provider: str = "claude") -> dict:
    """File -> full structured RFP JSON. This is the Module 1 deliverable."""
    text = extract_text(path)
    rfp = {
        "rfp_id": rfp_id or f"RFP-{uuid.uuid4().hex[:8]}",
        "filename": os.path.basename(path),
        "meta": extract_meta(text),
        "raw_text_chars": len(text),
        "warnings": (["Document contains little or no extractable text — "
                      "likely a scanned PDF. OCR required."] if len(text) < 200 else []),
        "extraction_method": "heuristic",
        "chunks": chunk_text(text),
        "requirements": extract_requirements(text),
        "evaluation_criteria": extract_criteria(text),
        "deadlines": extract_deadlines(text),
        "questions": extract_questions(text),
        "financials": extract_financials(text),
    }
    return llm_enrich(text, rfp, provider=provider) if use_llm else rfp


if __name__ == "__main__":
    import json
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else "sample_rfp.txt"
    rfp = parse_rfp(path, rfp_id="RFP-SAMPLE")
    print(json.dumps({k: (v if not isinstance(v, list) else f"{len(v)} items")
                      for k, v in rfp.items()}, indent=2))
    with open("sample_rfp_extracted.json", "w", encoding="utf-8") as f:
        json.dump(rfp, f, indent=2)
    print("Full JSON -> sample_rfp_extracted.json")
